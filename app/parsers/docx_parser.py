# app/parsers/docx_parser.py
from typing import List, Dict, Any
from docx import Document

def _cell_text(cell) -> str:
    # Join all paragraph runs inside the cell; strip control whitespace
    txt = "\n".join(p.text or "" for p in cell.paragraphs)
    return " ".join(txt.replace("\xa0", " ").split()).strip()

def _table_to_rows(tbl) -> List[List[str]]:
    rows: List[List[str]] = []
    for row in tbl.rows:
        rows.append([_cell_text(c) for c in row.cells])
    return rows

def parse_docx(path: str) -> Dict[str, Any]:
    """
    Return a uniform structure:
      {
        "raw_text": "<all visible text>",
        "tables": [ [ [cell, ...], ... ], ... ]
      }
    """
    doc = Document(path)

    # Collect raw text (paragraphs + table cell text)
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    tables: List[List[List[str]]] = []
    for t in doc.tables:
        rows = _table_to_rows(t)
        tables.append(rows)
        for r in rows:
            for c in r:
                if c:
                    parts.append(c)

    raw_text = "\n".join(s for s in parts if s).strip()
    return {"raw_text": raw_text, "tables": tables}
